#!/usr/bin/env python3

from pathlib import Path
import re
import shutil
import subprocess
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "template" / "Revised 983.docx"
CONTENT = ROOT / "content.md"
REFERENCES = ROOT / "references.md"
OUTPUT = ROOT / "output" / "vgadash_semester_paper_draft.docx"
OUTPUT_PDF = ROOT / "output" / "vgadash_semester_paper_draft.pdf"
ASSETS = ROOT / "assets"


def cleanup_output_dir() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    for pattern in ("~$*", "*_updated.docx", "*_updated.pdf"):
        for path in OUTPUT.parent.glob(pattern):
            try:
                path.unlink(missing_ok=True)
            except PermissionError:
                pass


def clear_document(doc: Document) -> None:
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def add_centered_paragraph(doc: Document, text: str, size: int, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_styled_paragraph(doc: Document, style: str, text: str) -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    return p


def set_run_font(paragraph, size=10):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)


def add_abstract(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Abstract")
    label = p.add_run("Abstract: ")
    label.bold = True
    label.font.name = "Times New Roman"
    label.font.size = Pt(10)
    body = p.add_run(text)
    body.font.name = "Times New Roman"
    body.font.size = Pt(10)


def add_keywords(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    label = p.add_run("Keywords: ")
    label.bold = True
    label.font.name = "Times New Roman"
    label.font.size = Pt(10)
    body = p.add_run(text)
    body.font.name = "Times New Roman"
    body.font.size = Pt(10)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="BodyText")
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading(doc: Document, text: str, level: int) -> None:
    style = {1: "Heading1", 2: "Heading2", 3: "Heading3"}[level]
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt({1: 12, 2: 11, 3: 10}[level])
    if level == 1:
        run.bold = True


def add_figure(doc: Document, name: str, caption: str) -> None:
    image_path = ASSETS / f"{name}.png"
    with Image.open(image_path) as img:
        width_px, height_px = img.size

    max_width = 3.15
    max_height = 3.8
    aspect = width_px / height_px
    width = max_width
    height = width / aspect
    if height > max_height:
        height = max_height
        width = height * aspect

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width), height=Inches(height))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = c.add_run(caption)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)


def add_comparison_table(doc: Document) -> None:
    doc.add_paragraph()
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run("Table 1. Positioning VGADASH against existing kernel debugging approaches.")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)

    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    hdr = table.rows[0].cells
    headers = ["Tool", "Works without userspace", "Needs network", "On-screen visibility", "Best fit"]
    for cell, text in zip(hdr, headers):
        cell.text = text

    rows = [
        ("VGADASH", "Yes", "No", "Yes", "Alive-but-unusable systems needing local visibility"),
        ("SysRq", "Yes", "No", "No", "Emergency commands without integrated display"),
        ("netconsole", "Yes", "Yes", "No", "Remote log export when networking is healthy"),
        ("pstore / ramoops", "Partly", "No", "No", "Post-reboot crash evidence"),
        ("kdump", "Partly", "No", "No", "Post-mortem crash capture"),
        ("earlyprintk", "Yes", "No", "Sometimes", "Very early boot logging only"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text

    widths = [0.72, 0.85, 0.6, 0.72, 1.55]
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)


def parse_content():
    lines = CONTENT.read_text(encoding="utf-8").splitlines()
    meta = {}
    sections = []
    current = None
    in_meta = True

    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if in_meta and ":" in line and not line.startswith("#"):
            key, value = line.split(":", 1)
            meta[key.strip()] = value.strip()
            i += 1
            continue
        if line.startswith("## "):
            in_meta = False
            current = {"level": 1, "title": line[3:].strip(), "blocks": []}
            sections.append(current)
            i += 1
            continue
        if line.startswith("### "):
            in_meta = False
            current = {"level": 2, "title": line[4:].strip(), "blocks": []}
            sections.append(current)
            i += 1
            continue
        if line.startswith("#### "):
            in_meta = False
            current = {"level": 3, "title": line[5:].strip(), "blocks": []}
            sections.append(current)
            i += 1
            continue

        if current is None:
            i += 1
            continue

        if line.startswith("[[FIGURE:"):
            match = re.match(r"\[\[FIGURE:([^|]+)\|(.+)\]\]", line)
            if match:
                current["blocks"].append(("figure", match.group(1), match.group(2)))
            i += 1
            continue

        if line.startswith("[[TABLE:comparison]]"):
            current["blocks"].append(("table", "comparison"))
            i += 1
            continue

        para_lines = [line.strip()]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not lines[i].startswith("[["):
            para_lines.append(lines[i].strip())
            i += 1
        current["blocks"].append(("paragraph", " ".join(para_lines)))

    return meta, sections


def load_references():
    refs = []
    for line in REFERENCES.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line:
            refs.append(line)
    return refs


def keep_section_properties(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = section.top_margin


def configure_double_column_section(section) -> None:
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols_el = cols[0] if cols else OxmlElement("w:cols")
    cols_el.set(qn("w:num"), "2")
    cols_el.set(qn("w:space"), "720")
    if not cols:
        sect_pr.append(cols_el)


def insert_double_column_section(doc: Document):
    section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    configure_double_column_section(section)
    return section


def append_page_number_run(paragraph) -> None:
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def add_centered_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        for p in list(footer.paragraphs):
            if p.text.strip():
                p.clear()
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.clear()
        append_page_number_run(paragraph)


def add_reference_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="references")
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def generate_pdf_with_docker() -> bool:
    docker = shutil.which("docker")
    if not docker:
        return False

    mount_root = ROOT.resolve().parents[1]
    docx_in_container = Path("/work") / OUTPUT.relative_to(mount_root)
    outdir_in_container = str((Path("/work") / OUTPUT.parent.relative_to(mount_root)).parent if False else Path("/work") / OUTPUT.parent.relative_to(mount_root))

    cmd = [
        docker,
        "run",
        "--rm",
        "-v",
        f"{mount_root}:/work",
        "-w",
        "/work",
        "ubuntu:22.04",
        "bash",
        "-lc",
        (
            "apt-get update >/dev/null && "
            "DEBIAN_FRONTEND=noninteractive apt-get install -y libreoffice-writer >/dev/null && "
            f"soffice --headless --convert-to pdf --outdir {outdir_in_container} "
            f"{docx_in_container} >/dev/null"
        ),
    ]
    subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    return OUTPUT_PDF.exists()


def main():
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    keep_section_properties(doc)

    meta, sections = parse_content()
    refs = load_references()

    add_centered_paragraph(doc, meta["Title"], 14, bold=True)
    add_centered_paragraph(doc, meta["Authors"], 12, bold=False)
    add_centered_paragraph(doc, meta["Institute"], 11, bold=False)
    add_centered_paragraph(doc, meta["Emails"], 10, bold=False)

    body_section_started = False
    first = True
    for section in sections:
        title = section["title"]
        if title == "Abstract":
            text = " ".join(block[1] for block in section["blocks"] if block[0] == "paragraph")
            add_abstract(doc, text)
            add_keywords(doc, meta["Keywords"])
            insert_double_column_section(doc)
            body_section_started = True
            first = False
            continue

        if not body_section_started:
            insert_double_column_section(doc)
            body_section_started = True

        add_heading(doc, title, section["level"])
        for block in section["blocks"]:
            if block[0] == "paragraph":
                add_body_paragraph(doc, block[1])
            elif block[0] == "figure":
                add_figure(doc, block[1], block[2])
            elif block[0] == "table" and block[1] == "comparison":
                add_comparison_table(doc)
        first = False

    add_heading(doc, "References", 1)
    for ref in refs:
        add_reference_paragraph(doc, ref)

    add_centered_page_numbers(doc)

    cleanup_output_dir()
    if OUTPUT.exists():
        try:
            OUTPUT.unlink()
        except PermissionError as exc:
            raise RuntimeError(f"Cannot overwrite {OUTPUT}; close it in Word and rerun.") from exc

    if OUTPUT_PDF.exists():
        OUTPUT_PDF.unlink()

    doc.save(str(OUTPUT))

    soffice = shutil.which("soffice")
    if soffice:
        subprocess.run(
            [
                soffice,
                "-env:UserInstallation=file:///tmp/lo_profile_vgadash",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(OUTPUT.parent),
                str(OUTPUT),
            ],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        generated_pdf = OUTPUT.with_suffix(".pdf")
        if generated_pdf != OUTPUT_PDF and generated_pdf.exists():
            if OUTPUT_PDF.exists():
                OUTPUT_PDF.unlink()
            generated_pdf.rename(OUTPUT_PDF)
    else:
        if not generate_pdf_with_docker():
            print("PDF not generated: neither soffice nor Docker conversion was available.")


if __name__ == "__main__":
    main()
